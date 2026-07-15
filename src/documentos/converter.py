"""
documentos/converter.py

All format-conversion logic for the "Converter Formato" sub-tab: images,
PDFs, DOCX and TXT files converted locally, no internet connection
required. This is a separate, document-focused converter from the
top-level converter.py (which handles video/audio/image conversion via
ffmpeg for the Downloads workflow) - they don't share code on purpose,
matching how downloader.py and the top-level converter.py already keep
their own small private helpers instead of sharing a base class.
"""

import os
import shutil
import subprocess
import tempfile

from PIL import Image
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from utils import find_poppler_bin_dir, no_window_flags, safe_filename, unique_path

IMAGE_EXTS = ["jpg", "jpeg", "png", "bmp", "webp", "tiff"]

POPPLER_INSTALL_MESSAGE = (
    "O poppler não foi encontrado neste computador. Ele é necessário para "
    "processar arquivos PDF (conversão de/para PDF).\n\n"
    "Windows: baixe em "
    "https://github.com/oschwartz10612/poppler-windows/releases, extraia o "
    ".zip e adicione a pasta \"Library\\bin\" ao PATH do Windows."
)


def wrap_poppler_error(exc: Exception) -> Exception:
    """Normalize a pdf2image failure into a clear PT-BR message when it's
    caused by poppler being missing."""
    message = str(exc)
    if "poppler" in message.lower() or exc.__class__.__name__ == "PDFInfoNotInstalledError":
        return RuntimeError(POPPLER_INSTALL_MESSAGE)
    return exc

# Explicit conversion table, matching the requested matrix - each format
# lists only the destinations that actually have an implementation below.
# It doesn't include its own format (e.g. "pdf" doesn't list "pdf"): staying
# in the same format is handled separately, as a passthrough/copy, rather
# than appearing as a normal "conversion".
CONVERSION_MATRIX = {
    "jpg": ["pdf", "png", "bmp", "webp", "tiff"],
    "jpeg": ["pdf", "png", "bmp", "webp", "tiff"],
    "png": ["pdf", "jpg", "bmp", "webp", "tiff"],
    "bmp": ["pdf", "jpg", "png", "webp", "tiff"],
    "webp": ["pdf", "jpg", "png", "bmp", "tiff"],
    "tiff": ["pdf", "jpg", "png", "bmp", "webp"],
    "pdf": ["jpg", "png", "txt", "docx"],
    "docx": ["pdf", "txt"],
    "txt": ["pdf", "docx"],
}


def detect_format(file_path: str):
    """Return (extension, category). category is one of "image", "pdf",
    "docx", "txt", or None if unsupported."""
    ext = os.path.splitext(file_path)[1].lower().lstrip(".")
    if ext in IMAGE_EXTS:
        return ext, "image"
    if ext == "pdf":
        return ext, "pdf"
    if ext == "docx":
        return ext, "docx"
    if ext == "txt":
        return ext, "txt"
    return ext, None


def available_targets(source_ext: str):
    return list(CONVERSION_MATRIX.get(source_ext.lower(), []))


def can_convert(source_ext: str, target_ext: str) -> bool:
    """Whether a specific file can reach a specific target format - either
    via a real conversion, or via the same-format passthrough copy."""
    source_ext = source_ext.lower()
    target_ext = target_ext.lower()
    if source_ext == target_ext:
        return True
    return target_ext in available_targets(source_ext)


def convert_file(source_path: str, target_ext: str, output_dir: str, progress_cb=None):
    """Convert a single file, returning a list of output paths (usually one,
    except PDF->image which produces one file per page)."""
    ext = os.path.splitext(source_path)[1].lower().lstrip(".")
    target_ext = target_ext.lower()
    os.makedirs(output_dir, exist_ok=True)

    # Same-format passthrough: e.g. a PDF already in the list, converting
    # "to PDF" - just copy it instead of erroring on an unhandled pair.
    if ext == target_ext:
        return [_copy_passthrough(source_path, output_dir)]

    if ext in IMAGE_EXTS and target_ext == "pdf":
        return [_image_to_pdf(source_path, output_dir)]
    if ext in IMAGE_EXTS and target_ext in IMAGE_EXTS:
        return [_image_to_image(source_path, target_ext, output_dir)]
    if ext == "pdf" and target_ext in ("jpg", "png"):
        return _pdf_to_images(source_path, target_ext, output_dir, progress_cb)
    if ext == "pdf" and target_ext == "docx":
        return [_pdf_to_docx(source_path, output_dir)]
    if ext == "pdf" and target_ext == "txt":
        return [_pdf_to_txt(source_path, output_dir, progress_cb)]
    if ext == "docx" and target_ext == "pdf":
        return [_docx_to_pdf(source_path, output_dir)]
    if ext == "docx" and target_ext == "txt":
        return [_docx_to_txt(source_path, output_dir)]
    if ext == "txt" and target_ext == "pdf":
        return [_txt_to_pdf(source_path, output_dir)]
    if ext == "txt" and target_ext == "docx":
        return [_txt_to_docx(source_path, output_dir)]

    raise ValueError(f"Conversão de .{ext} para .{target_ext} não é suportada.")


def merge_to_pdf(paths, output_dir: str, output_name: str = None, passwords: dict = None,
                  password_callback=None) -> str:
    """Merge several files (images, PDFs, DOCX, TXT - any mix) into a
    single PDF, one source file's pages appended after another, in the
    given order. Returns the output path.

    Every PDF (encrypted or not) is opened here, inside this function - so
    this must be called off the GUI thread, since parsing a PDF's
    structure can take a noticeable while for large/unusual files.

    `passwords` is an optional {path: password} map with passwords already
    known upfront. `password_callback`, if given, is called as
    `password_callback(path)` for any encrypted PDF not already covered by
    `passwords`, and must return the password to try (or None to skip
    decryption). This indirection lets the caller (typically a background
    worker) ask the GUI thread for a password on demand, without this
    function needing to know anything about Qt."""
    from pypdf import PdfReader, PdfWriter

    if not paths:
        raise ValueError("Nenhum arquivo para mesclar.")
    passwords = passwords or {}

    os.makedirs(output_dir, exist_ok=True)
    writer = PdfWriter()
    temp_dir = tempfile.mkdtemp(prefix="docmerge_")
    try:
        for path in paths:
            ext = os.path.splitext(path)[1].lower().lstrip(".")
            if ext == "pdf":
                # Password-protected PDFs are decrypted before being
                # appended - without this, pypdf refuses to read the pages
                # of an encrypted PDF.
                reader = PdfReader(path)
                if reader.is_encrypted:
                    password = passwords.get(path)
                    if password is None and password_callback is not None:
                        password = password_callback(path)
                    if password is None or reader.decrypt(password) == 0:
                        raise ValueError(
                            f"Senha incorreta (ou não informada) para o arquivo protegido "
                            f"\"{os.path.basename(path)}\"."
                        )
                writer.append(reader)
            else:
                pdf_path = convert_file(path, "pdf", temp_dir)[0]
                writer.append(pdf_path)

        base_name = safe_filename(os.path.splitext(output_name)[0]) if output_name else "documento_mesclado"
        out_path = unique_path(output_dir, base_name, "pdf")
        with open(out_path, "wb") as f:
            writer.write(f)
    finally:
        writer.close()
        shutil.rmtree(temp_dir, ignore_errors=True)

    return out_path


# ---------------------------------------------------------------------------
# Individual conversions
# ---------------------------------------------------------------------------

def _copy_passthrough(path: str, output_dir: str) -> str:
    """Same-format "conversion": just copy the file to the output folder
    under a non-colliding name."""
    base = safe_filename(os.path.splitext(os.path.basename(path))[0])
    ext = os.path.splitext(path)[1].lstrip(".")
    out_path = unique_path(output_dir, base, ext)
    shutil.copy2(path, out_path)
    return out_path


def _image_to_pdf(path: str, output_dir: str) -> str:
    image = Image.open(path).convert("RGB")
    page_w, page_h = A4
    draw_w, draw_h, x, y = _fit_to_page(image.size, (page_w, page_h))

    base = safe_filename(os.path.splitext(os.path.basename(path))[0])
    out_path = unique_path(output_dir, base, "pdf")
    c = canvas.Canvas(out_path, pagesize=A4)
    c.drawImage(ImageReader(image), x, y, width=draw_w, height=draw_h)
    c.save()
    return out_path


def _image_to_image(path: str, target_ext: str, output_dir: str) -> str:
    image = Image.open(path)
    if target_ext in ("jpg", "jpeg") and image.mode in ("RGBA", "P"):
        image = image.convert("RGB")
    save_format = "JPEG" if target_ext in ("jpg", "jpeg") else target_ext.upper()

    base = safe_filename(os.path.splitext(os.path.basename(path))[0])
    out_path = unique_path(output_dir, base, target_ext)
    image.save(out_path, save_format)
    return out_path


def _pdf_to_images(path: str, target_ext: str, output_dir: str, progress_cb=None):
    from pdf2image import convert_from_path

    try:
        # Explicit poppler_path: MasterApp.bat downloads Poppler into
        # tools/poppler instead of registering it on PATH, so pdf2image
        # needs to be told exactly where to find it. Passing None (when
        # nothing is found - e.g. poppler already IS on PATH) keeps the
        # previous behavior intact.
        pages = convert_from_path(path, dpi=200, poppler_path=find_poppler_bin_dir() or None)
    except Exception as exc:  # noqa: BLE001 - normalize to a clear PT-BR message
        raise wrap_poppler_error(exc)

    save_format = "JPEG" if target_ext in ("jpg", "jpeg") else target_ext.upper()
    base = safe_filename(os.path.splitext(os.path.basename(path))[0])
    out_paths = []
    total = len(pages)
    for index, page in enumerate(pages, start=1):
        out_path = unique_path(output_dir, f"{base}_pg{index}", target_ext)
        if target_ext in ("jpg", "jpeg"):
            page = page.convert("RGB")
        page.save(out_path, save_format)
        out_paths.append(out_path)
        if progress_cb:
            progress_cb(index, total)
    return out_paths


def _pdf_to_docx(path: str, output_dir: str) -> str:
    from pdf2docx import Converter

    base = safe_filename(os.path.splitext(os.path.basename(path))[0])
    out_path = unique_path(output_dir, base, "docx")
    converter = Converter(path)
    try:
        converter.convert(out_path)
    finally:
        converter.close()
    return out_path


def _pdf_to_txt(path: str, output_dir: str, progress_cb=None) -> str:
    import pdfplumber

    base = safe_filename(os.path.splitext(os.path.basename(path))[0])
    out_path = unique_path(output_dir, base, "txt")
    lines = []
    with pdfplumber.open(path) as pdf:
        total = len(pdf.pages)
        for index, page in enumerate(pdf.pages, start=1):
            lines.append(page.extract_text() or "")
            if progress_cb:
                progress_cb(index, total)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(lines))
    return out_path


def _docx_to_pdf(path: str, output_dir: str) -> str:
    base = safe_filename(os.path.splitext(os.path.basename(path))[0])
    out_path = unique_path(output_dir, base, "pdf")

    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(path, out_path)
        if os.path.exists(out_path):
            return out_path
    except Exception:
        pass  # fall through to the LibreOffice fallback below

    libreoffice = shutil.which("soffice") or shutil.which("libreoffice")
    if libreoffice:
        subprocess.run(
            [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", output_dir, path],
            capture_output=True,
            timeout=120,
            creationflags=no_window_flags(),
        )
        if os.path.exists(out_path):
            return out_path

    raise RuntimeError(
        "Não foi possível converter DOCX para PDF: é necessário ter o Microsoft "
        "Word instalado (Windows) ou o LibreOffice instalado no sistema."
    )


def _docx_to_txt(path: str, output_dir: str) -> str:
    from docx import Document

    base = safe_filename(os.path.splitext(os.path.basename(path))[0])
    out_path = unique_path(output_dir, base, "txt")
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)
    return out_path


def _txt_to_pdf(path: str, output_dir: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    base = safe_filename(os.path.splitext(os.path.basename(path))[0])
    return _text_to_pdf(text, output_dir, base)


def _txt_to_docx(path: str, output_dir: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    base = safe_filename(os.path.splitext(os.path.basename(path))[0])
    return _text_to_docx(text, output_dir, base)


def _text_to_docx(text: str, output_dir: str, base_name: str) -> str:
    from docx import Document

    os.makedirs(output_dir, exist_ok=True)
    out_path = unique_path(output_dir, safe_filename(base_name), "docx")
    document = Document()
    for paragraph in text.split("\n"):
        document.add_paragraph(paragraph)
    document.save(out_path)
    return out_path


def _text_to_pdf(text: str, output_dir: str, base_name: str) -> str:
    """Render the text as real, selectable PDF text (not an image), which
    makes it inherently searchable."""
    os.makedirs(output_dir, exist_ok=True)
    out_path = unique_path(output_dir, safe_filename(base_name), "pdf")
    page_w, page_h = A4
    margin = 40
    line_height = 14

    c = canvas.Canvas(out_path, pagesize=A4)
    c.setFont("Helvetica", 10)
    y = page_h - margin
    for raw_line in text.split("\n"):
        for line in _wrap_text_line(raw_line, 95):
            if y < margin:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = page_h - margin
            c.drawString(margin, y, line)
            y -= line_height
    c.save()
    return out_path


def _wrap_text_line(line: str, width: int):
    if not line:
        return [""]
    words = line.split(" ")
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) > width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines or [""]


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------

def _fit_to_page(image_size, page_size):
    """Scale an image to fit centered inside a page, preserving aspect
    ratio. Returns (draw_width, draw_height, x, y)."""
    img_w, img_h = image_size
    page_w, page_h = page_size
    scale = min(page_w / img_w, page_h / img_h)
    draw_w, draw_h = img_w * scale, img_h * scale
    x = (page_w - draw_w) / 2
    y = (page_h - draw_h) / 2
    return draw_w, draw_h, x, y
